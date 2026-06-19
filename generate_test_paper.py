import os
import re
import sys
import random
import openpyxl
from openpyxl.styles import Font as xlFont, PatternFill, Alignment as xlAlignment
import docx
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Ensure utf-8 terminal print support
sys.stdout.reconfigure(encoding='utf-8')

# Helper to add custom XML borders to table cells
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

# Smart splitting logic for bilingual question/option text
def split_bilingual(text):
    if not text:
        return "", ""
    text = str(text).strip()
    
    # Case 1: Split by double newline + bracket e.g. "Hindi \n\n[English]"
    match = re.search(r'^(.*?)\s*\n+\s*\[(.*?)\]\s*$', text, re.DOTALL)
    if match:
        return match.group(1).strip(), match.group(2).strip()
        
    # Case 2: Split by bracket at the end e.g. "Hindi [English]"
    if "[" in text and text.endswith("]"):
        parts = text.split("[")
        if len(parts) >= 2:
            hindi = "".join(parts[:-1]).strip()
            english = parts[-1].rstrip("]").strip()
            return hindi, english
            
    # Case 3: Split by " / " divider
    if " / " in text:
        parts = text.split(" / ")
        return parts[0].strip(), parts[1].strip()
        
    # Case 4: Split by "/" divider
    if "/" in text:
        parts = text.split("/")
        return parts[0].strip(), parts[1].strip()
        
    return text, ""

# Dynamic page count field injector
def add_num_pages_field(paragraph):
    r1 = paragraph.add_run("पुस्तिका में पृष्ठों की संख्या : ")
    r1.font.name = 'Arial'
    r1.font.size = Pt(9.5)
    
    run_field1 = paragraph.add_run()
    run_field1.font.name = 'Arial'
    run_field1.font.size = Pt(9.5)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run_field1._r.append(fldChar1)
    run_field1._r.append(instrText1)
    run_field1._r.append(fldChar2)
    run_field1._r.append(fldChar3)
    
    r2 = paragraph.add_run("\n[Number of Pages in Booklet : ")
    r2.font.name = 'Arial'
    r2.font.size = Pt(9.5)
    
    run_field2 = paragraph.add_run()
    run_field2.font.name = 'Arial'
    run_field2.font.size = Pt(9.5)
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    fldChar5 = OxmlElement('w:fldChar')
    fldChar5.set(qn('w:fldCharType'), 'separate')
    fldChar6 = OxmlElement('w:fldChar')
    fldChar6.set(qn('w:fldCharType'), 'end')
    
    run_field2._r.append(fldChar4)
    run_field2._r.append(instrText2)
    run_field2._r.append(fldChar5)
    run_field2._r.append(fldChar6)
    
    r3 = paragraph.add_run("]")
    r3.font.name = 'Arial'
    r3.font.size = Pt(9.5)

# Save separate Excel Answer Key
def save_excel_answer_key(questions, excel_key_path, booklet_no):
    print(f"Generating Excel Answer Key: {excel_key_path}")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Answer Key"
    
    # Metadata Row 1: Booklet Number
    ws.merge_cells('A1:B1')
    ws['A1'] = f"Booklet No: {booklet_no}"
    ws['A1'].font = xlFont(name="Arial", size=11, bold=True, color="1A237E")
    ws['A1'].alignment = xlAlignment(horizontal="center")
    
    # Row 2 Headers
    ws['A2'] = "Question No"
    ws['B2'] = "answer"
    
    # Styles for headers
    header_font = xlFont(name="Arial", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1A237E", end_color="1A237E", fill_type="solid")
    header_align = xlAlignment(horizontal="center", vertical="center")
    
    ws['A2'].font = header_font
    ws['A2'].fill = header_fill
    ws['A2'].alignment = header_align
    ws['B2'].font = header_font
    ws['B2'].fill = header_fill
    ws['B2'].alignment = header_align
    
    for idx, q in enumerate(questions, 3):
        ws.cell(row=idx, column=1, value=q['qno']).alignment = xlAlignment(horizontal="center")
        ws.cell(row=idx, column=2, value=str(q['answer']) if q['answer'] else "").alignment = xlAlignment(horizontal="center")
        
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 15
    
    wb.save(excel_key_path)
    print(f"[SUCCESS] Excel Answer Key saved at: {excel_key_path}")

def generate_test_from_excel(excel_path, docx_path, sheet_name=None, topic_filter=None):
    print(f"Reading questions from Excel: {excel_path}")
    if not os.path.exists(excel_path):
        print(f"[ERROR] Excel file not found at: {excel_path}")
        return
        
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    if sheet_name is None:
        sheet_name = next((s for s in wb.sheetnames if "PYQ Questions" in s), wb.sheetnames[0])
    
    ws = wb[sheet_name]
    print(f"Using sheet: {sheet_name}")
    if topic_filter:
        print(f"Applying Topic/Sub-Topic filter: '{topic_filter}'")
    
    # Generate unique 6-digit booklet number
    booklet_no = str(random.randint(100000, 999999))
    
    # Identify column indices
    q_col, opt_start, ans_col = 3, 4, 8 # Defaults based on standard format
    
    # Read rows (skipping metadata rows)
    questions = []
    for r_idx, row in enumerate(ws.iter_rows(min_row=4, values_only=True), 4):
        if not row or len(row) < 9: continue
        q_val = row[q_col]
        if not q_val: continue
        
        # Topic Filter check
        if topic_filter:
            topic_val = str(row[1]) if row[1] is not None else ""
            subtopic_val = str(row[2]) if row[2] is not None else ""
            if topic_filter.lower() not in topic_val.lower() and topic_filter.lower() not in subtopic_val.lower():
                continue
        
        opts = [row[opt_start], row[opt_start+1], row[opt_start+2], row[opt_start+3]]
        ans_val = row[ans_col]
        
        questions.append({
            'qno': len(questions) + 1,
            'question': str(q_val),
            'options': [str(o) if o is not None else "" for o in opts],
            'answer': ans_val if ans_val is not None else ""
        })

    print(f"Parsed {len(questions)} questions.")
    
    # Build docx Document
    doc = docx.Document()
    
    # 1. Cover Page Setup (Section 1)
    cover_section = doc.sections[0]
    cover_section.top_margin = Inches(0.6)
    cover_section.bottom_margin = Inches(0.6)
    cover_section.left_margin = Inches(0.6)
    cover_section.right_margin = Inches(0.6)
    
    # RPSC Header on cover page (Separate paragraphs in Arial for precise spacing control)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(1)
    p_title.paragraph_format.line_spacing = 1.0
    run_h1 = p_title.add_run("राजस्थान प्रतियोगी परीक्षा")
    run_h1.font.name = 'Arial'
    run_h1.font.size = Pt(22)
    run_h1.font.bold = True
    
    # Date box paragraph
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(1)
    p_date.paragraph_format.space_after = Pt(6)
    p_date.paragraph_format.line_spacing = 1.0
    run_h2 = p_date.add_run("परीक्षा दिनांक - [                   ]")
    run_h2.font.name = 'Arial'
    run_h2.font.size = Pt(14)
    run_h2.font.bold = True
    
    # Table for Booklet metadata
    table_meta = doc.add_table(rows=3, cols=3)
    table_meta.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    col_widths = [Inches(2.3), Inches(2.7), Inches(2.3)]
    for r in table_meta.rows:
        for idx, w in enumerate(col_widths):
            r.cells[idx].width = w

    # Metadata cells content (Styled in Arial with tight padding)
    # Cell 0, 0: Dynamic Page Count using Word field
    cell00 = table_meta.cell(0, 0).paragraphs[0]
    cell00.paragraph_format.space_before = Pt(1)
    cell00.paragraph_format.space_after = Pt(1)
    cell00.paragraph_format.line_spacing = 1.15
    add_num_pages_field(cell00)
    
    c01 = table_meta.cell(0, 1).paragraphs[0]
    c01.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c01.paragraph_format.space_before = Pt(2)
    c01.paragraph_format.space_after = Pt(2)
    r_asr = c01.add_run("ASR-24")
    r_asr.font.name = 'Arial'
    r_asr.font.size = Pt(22)
    r_asr.font.bold = True
    
    # Cell 0, 2: Booklet Number and simulated Barcode (No duplicate newlines)
    cell02 = table_meta.cell(0, 2).paragraphs[0]
    cell02.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell02.paragraph_format.space_before = Pt(1)
    cell02.paragraph_format.space_after = Pt(1)
    cell02.paragraph_format.line_spacing = 1.15
    r02_title = cell02.add_run("प्रश्न-पुस्तिका संख्या व बारकोड\n[Question Booklet No. & Barcode]")
    r02_title.font.name = 'Arial'
    r02_title.font.size = Pt(9.5)
    r02_num = cell02.add_run(f"\n{booklet_no}")
    r02_num.font.name = 'Arial'
    r02_num.font.size = Pt(14)
    r02_num.font.bold = True
    r02_bar = cell02.add_run(f"\n||||||| | ||||| | ||| ||")
    r02_bar.font.name = 'Arial'
    r02_bar.font.size = Pt(12)

    # Cell 1, 0: Dynamic Question Count
    cell10 = table_meta.cell(1, 0).paragraphs[0]
    cell10.paragraph_format.space_before = Pt(1)
    cell10.paragraph_format.space_after = Pt(1)
    cell10.paragraph_format.line_spacing = 1.15
    r10 = cell10.add_run(f"पुस्तिका में प्रश्नों की संख्या : {len(questions)}\n[No. of Questions in Booklet : {len(questions)}]")
    r10.font.name = 'Arial'
    r10.font.size = Pt(9.5)
    
    c11 = table_meta.cell(1, 1).paragraphs[0]
    c11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c11.paragraph_format.space_before = Pt(1)
    c11.paragraph_format.space_after = Pt(1)
    c11.paragraph_format.line_spacing = 1.15
    r11 = c11.add_run("इस प्रश्न-पुस्तिका को तब तक न खोलें जब तक कहा न जाए\n[Do not open this Booklet until asked]")
    r11.font.name = 'Arial'
    r11.font.size = Pt(9)
    
    # Cell 1, 2: Empty Paper Code
    cell12 = table_meta.cell(1, 2).paragraphs[0]
    cell12.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell12.paragraph_format.space_before = Pt(1)
    cell12.paragraph_format.space_after = Pt(1)
    cell12.paragraph_format.line_spacing = 1.15
    r12 = cell12.add_run("Paper Code : [       ]")
    r12.font.name = 'Arial'
    r12.font.size = Pt(11)
    
    cell20 = table_meta.cell(2, 0).paragraphs[0]
    cell20.paragraph_format.space_before = Pt(1)
    cell20.paragraph_format.space_after = Pt(1)
    cell20.paragraph_format.line_spacing = 1.15
    r20 = cell20.add_run("समय : 03 घण्टे + 10 मिनट अतिरिक्त*\n[Time : 03 Hours + 10 Minutes Extra*]")
    r20.font.name = 'Arial'
    r20.font.size = Pt(9.5)
    
    c21 = table_meta.cell(2, 1).paragraphs[0]
    c21.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c21.paragraph_format.space_before = Pt(1)
    c21.paragraph_format.space_after = Pt(1)
    r_sub = c21.add_run("Sub : G.K. & G.S.")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    
    cell22 = table_meta.cell(2, 2).paragraphs[0]
    cell22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell22.paragraph_format.space_before = Pt(1)
    cell22.paragraph_format.space_after = Pt(1)
    cell22.paragraph_format.line_spacing = 1.15
    r22 = cell22.add_run("अधिकतम अंक : 200\n[Maximum Marks : 200]")
    r22.font.name = 'Arial'
    r22.font.size = Pt(9.5)

    # Style metadata cell borders
    border_style = {"sz": 6, "val": "single", "color": "000000"}
    for row in table_meta.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
            
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(1)
    p_spacer.paragraph_format.space_after = Pt(1)
    p_spacer.paragraph_format.line_spacing = 1.0
    run_sp = p_spacer.add_run()
    run_sp.font.size = Pt(1)

    # Cover Page Instructions table (Bilingual 2-Column)
    table_inst = doc.add_table(rows=1, cols=2)
    table_inst.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table_inst.autofit = False
    table_inst.columns[0].width = Inches(3.6)
    table_inst.columns[1].width = Inches(3.6)
    
    # Left Cell: Hindi Instructions
    cell_hin = table_inst.cell(0, 0)
    p_hin = cell_hin.paragraphs[0]
    p_hin.paragraph_format.space_before = Pt(4)
    p_hin.paragraph_format.space_after = Pt(4)
    p_hin.paragraph_format.line_spacing = 1.15
    r_hi = p_hin.add_run("परीक्षार्थियों के लिए निर्देश")
    r_hi.font.name = 'Arial'
    r_hi.font.bold = True
    r_hi.font.size = Pt(11.5)
    p_hin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    instructions_hin = [
        "प्रत्येक प्रश्न के लिये एक विकल्प भरना अनिवार्य है।",
        "सभी प्रश्नों के अंक समान हैं।",
        "प्रत्येक प्रश्न का मात्र एक ही उत्तर दीजिए। एक से अधिक उत्तर देने की दशा में प्रश्न के उत्तर को गलत माना जाएगा।",
        "OMR उत्तर-पत्रक इस प्रश्न-पुस्तिका के अन्दर रखा है। जब आपको प्रश्न-पुस्तिका खोलने को कहा जाए, तो उत्तर-पत्रक निकालकर ध्यान से केवल नीले बॉल पॉइंट पेन से विवरण भरें।",
        "कृपया अपना रोल नम्बर ओ.एम.आर. उत्तर-पत्रक पर सावधानीपूर्वक सही भरें। गलत रोल नम्बर भरने पर परीक्षार्थी स्वयं उत्तरदायी होगा।",
        "ओ.एम.आर. उत्तर-पत्रक में करेक्शन पेन/व्हाइटनर/सफेदा का उपयोग निषिद्ध है।",
        "प्रत्येक गलत उत्तर के लिए प्रश्न अंक का 1/3 भाग काटा जायेगा। गलत उत्तर से तात्पर्य अशुद्ध उत्तर अथवा किसी भी प्रश्न के एक से अधिक उत्तर से है।",
        "प्रत्येक प्रश्न के पाँच विकल्प दिये गये हैं, जिन्हें क्रमशः 1, 2, 3, 4, 5 अंकित किया गया है। अभ्यर्थी को सही उत्तर निर्दिष्ट करते हुए उनमें से केवल एक गोले (बबल) को उत्तर-पत्रक पर नीले बॉल पॉइंट पेन से गहरा करना है।",
        "यदि आप प्रश्न का उत्तर नहीं देना चाहते हैं तो उत्तर-पत्रक में पाँचवें (5) विकल्प को गहरा करें। यदि पाँच में से कोई भी गोला गहरा नहीं किया जाता है, तो ऐसे प्रश्न के लिये प्रश्न अंक का 1/3 भाग काटा जायेगा।",
        "प्रश्न-पत्र हल करने के उपरांत अभ्यर्थी अनिवार्य रूप से ओ.एम.आर. उत्तर-पत्रक जांच लें कि समस्त प्रश्नों के लिये एक विकल्प (गोला) भर दिया गया है। इसके लिये निर्धारित समय से 10 मिनट का अतिरिक्त समय दिया गया है।"
    ]
    
    for idx, text in enumerate(instructions_hin, 1):
        p = cell_hin.add_paragraph()
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after = Pt(0.5)
        p.paragraph_format.line_spacing = 1.15
        r_num = p.add_run(f"{idx}. ")
        r_num.font.name = 'Arial'
        r_num.font.bold = True
        r_num.font.size = Pt(8.5)
        r_txt = p.add_run(text)
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(8.5)

    # Right Cell: English Instructions
    cell_eng = table_inst.cell(0, 1)
    p_eng = cell_eng.paragraphs[0]
    p_eng.paragraph_format.space_before = Pt(4)
    p_eng.paragraph_format.space_after = Pt(4)
    p_eng.paragraph_format.line_spacing = 1.15
    r_en = p_eng.add_run("INSTRUCTIONS FOR CANDIDATES")
    r_en.font.name = 'Arial'
    r_en.font.bold = True
    r_en.font.size = Pt(11.5)
    p_eng.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    instructions_eng = [
        "It is mandatory to fill one option for each question.",
        "All questions carry equal marks.",
        "Only one answer is to be given for each question. If more than one answers are marked, it would be treated as wrong answer.",
        "The OMR Answer Sheet is inside this Question Booklet. When you are directed to open the Question Booklet, take out the Answer Sheet and fill in the particulars carefully with Blue Ball Point Pen only.",
        "Please correctly fill your Roll Number in OMR Answer Sheet. Candidate will themself be responsible for filling wrong Roll No.",
        "Use of Correction Pen/Whitner in the OMR Answer Sheet is strictly forbidden.",
        "1/3 part of the mark(s) of each question will be deducted for each wrong answer. A wrong answer means an incorrect answer or more than one answers for any question.",
        "Each question has five options marked as 1, 2, 3, 4, 5. You have to darken only one circle (bubble) indicating the correct answer on the Answer Sheet using BLUE BALL POINT PEN.",
        "If you are not attempting a question then you have to darken the circle '5'. If none of the five circles is darkened, one third (1/3) part of the marks of question shall be deducted.",
        "After solving question paper, candidate must ascertain that he/she has darkened one of the circles (bubbles) for each of the questions. Extra time of 10 minutes beyond scheduled time, is provided for this."
    ]
    
    for idx, text in enumerate(instructions_eng, 1):
        p = cell_eng.add_paragraph()
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after = Pt(0.5)
        p.paragraph_format.line_spacing = 1.15
        r_num = p.add_run(f"{idx}. ")
        r_num.font.name = 'Arial'
        r_num.font.bold = True
        r_num.font.size = Pt(8.5)
        r_txt = p.add_run(text)
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(8.5)

    # Style instructions border lines
    set_cell_border(cell_hin, right={"sz": 6, "val": "single", "color": "000000"})
    set_cell_border(cell_eng, left={"sz": 6, "val": "single", "color": "000000"})
    for cell in (cell_hin, cell_eng):
        set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"}, bottom={"sz": 12, "val": "single", "color": "000000"})
    set_cell_border(cell_hin, left={"sz": 12, "val": "single", "color": "000000"})
    set_cell_border(cell_eng, right={"sz": 12, "val": "single", "color": "000000"})

    # 2. Inside Pages Section Setup (2-Column Hindi-only layout grid - Flowing next question in next column)
    inside_section = doc.add_section(WD_SECTION.NEW_PAGE)
    inside_section.top_margin = Inches(0.6)
    inside_section.bottom_margin = Inches(0.6)
    inside_section.left_margin = Inches(0.6)
    inside_section.right_margin = Inches(0.6)
    
    table_ques = doc.add_table(rows=0, cols=2)
    table_ques.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table_ques.autofit = False
    table_ques.columns[0].width = Inches(3.6)
    table_ques.columns[1].width = Inches(3.6)
    
    # Style separators
    border_middle = {"sz": 6, "val": "single", "color": "D3D3D3"} # light gray vertical line
    border_bottom = {"sz": 4, "val": "single", "color": "E0E0E0"} # subtle question divider
    
    # Loop over questions two at a time: left cell gets Q_N, right cell gets Q_N+1
    for idx in range(0, len(questions), 2):
        row = table_ques.add_row()
        row.cells[0].width = Inches(3.6)
        row.cells[1].width = Inches(3.6)
        
        q_left = questions[idx]
        q_right = questions[idx+1] if idx+1 < len(questions) else None
        
        # --- LEFT CELL (q_left) ---
        cell_l = row.cells[0]
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(6)
        p_l.paragraph_format.line_spacing = 1.15
        
        hindi_q_l, _ = split_bilingual(q_left['question'])
        opts_l = []
        for o in q_left['options']:
            ho, _ = split_bilingual(o)
            opts_l.append(ho)
        opts_l.append("अनुत्तरित प्रश्न") # 5th option
        
        r_num_l = p_l.add_run(f"{q_left['qno']}.  ")
        r_num_l.font.name = 'Times New Roman'
        r_num_l.font.bold = True
        r_num_l.font.size = Pt(11.5)
        
        r_txt_l = p_l.add_run(hindi_q_l)
        r_txt_l.font.name = 'Times New Roman'
        r_txt_l.font.size = Pt(11.5)
        
        for o_idx, o_txt in enumerate(opts_l, 1):
            r_opt = p_l.add_run(f"\n    ({o_idx}) {o_txt}")
            r_opt.font.name = 'Times New Roman'
            r_opt.font.size = Pt(10.5)
            
        set_cell_border(cell_l, right=border_middle, bottom=border_bottom)
        
        # --- RIGHT CELL (q_right) ---
        cell_r = row.cells[1]
        p_r = cell_r.paragraphs[0]
        p_r.paragraph_format.space_after = Pt(6)
        p_r.paragraph_format.line_spacing = 1.15
        
        if q_right:
            hindi_q_r, _ = split_bilingual(q_right['question'])
            opts_r = []
            for o in q_right['options']:
                ho, _ = split_bilingual(o)
                opts_r.append(ho)
            opts_r.append("अनुत्तरित प्रश्न") # 5th option
            
            r_num_r = p_r.add_run(f"{q_right['qno']}.  ")
            r_num_r.font.name = 'Times New Roman'
            r_num_r.font.bold = True
            r_num_r.font.size = Pt(11.5)
            
            r_txt_r = p_r.add_run(hindi_q_r)
            r_txt_r.font.name = 'Times New Roman'
            r_txt_r.font.size = Pt(11.5)
            
            for o_idx, o_txt in enumerate(opts_r, 1):
                r_opt = p_r.add_run(f"\n    ({o_idx}) {o_txt}")
                r_opt.font.name = 'Times New Roman'
                r_opt.font.size = Pt(10.5)
                
        set_cell_border(cell_r, left=border_middle, bottom=border_bottom)
        
    # 3. Add Answer Key Grid at the end of the Word document
    doc.add_page_break()
    p_ak_meta = doc.add_paragraph()
    p_ak_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ak_meta = p_ak_meta.add_run(f"प्रश्न-पुस्तिका संख्या / Question Booklet No. : {booklet_no}\n")
    r_ak_meta.font.name = 'Times New Roman'
    r_ak_meta.font.size = Pt(11)
    r_ak_meta.font.bold = True
    
    p_ak = doc.add_paragraph()
    p_ak.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ak = p_ak.add_run("उत्तर कुंजी / ANSWER KEY")
    r_ak.font.name = 'Times New Roman'
    r_ak.font.size = Pt(14)
    r_ak.font.bold = True
    
    # 5 pairs of (Q.No, Ans) = 10 columns
    cols_count = 10
    rows_count = (len(questions) + 4) // 5
    
    table_ak = doc.add_table(rows=rows_count + 1, cols=cols_count)
    table_ak.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table_ak.autofit = False
    
    ak_widths = [Inches(0.6), Inches(0.8)] * 5
    for r in table_ak.rows:
        for idx, w in enumerate(ak_widths):
            r.cells[idx].width = w
            
    # Headers
    for pair in range(5):
        cell_q = table_ak.cell(0, pair * 2)
        cell_a = table_ak.cell(0, pair * 2 + 1)
        
        cell_q.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        rq = cell_q.paragraphs[0].add_run("Q.No.")
        rq.font.name = 'Times New Roman'
        rq.font.size = Pt(9.5)
        rq.font.bold = True
        
        ra = cell_a.paragraphs[0].add_run("Ans")
        ra.font.name = 'Times New Roman'
        ra.font.size = Pt(9.5)
        ra.font.bold = True
        
        set_cell_border(cell_q, top=border_style, bottom=border_style, left=border_style, right=border_style)
        set_cell_border(cell_a, top=border_style, bottom=border_style, left=border_style, right=border_style)
        
    # Populate key cells and borders
    border_cell = {"sz": 4, "val": "single", "color": "D3D3D3"}
    for idx, q in enumerate(questions):
        col_pair = idx % 5
        row_idx = (idx // 5) + 1
        
        cell_q = table_ak.cell(row_idx, col_pair * 2)
        cell_a = table_ak.cell(row_idx, col_pair * 2 + 1)
        
        cell_q.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        rq = cell_q.paragraphs[0].add_run(str(q['qno']))
        rq.font.name = 'Times New Roman'
        rq.font.size = Pt(9.5)
        
        ans_str = f"({q['answer']})" if q['answer'] else "-"
        ra = cell_a.paragraphs[0].add_run(ans_str)
        ra.font.name = 'Times New Roman'
        ra.font.size = Pt(9.5)
        
    # Apply borders to all data rows
    for r_idx in range(1, rows_count + 1):
        for c_idx in range(cols_count):
            cell = table_ak.cell(r_idx, c_idx)
            set_cell_border(cell, top=border_cell, bottom=border_cell, left=border_cell, right=border_cell)
            
    doc.save(docx_path)
    print(f"[SUCCESS] Professional Test Paper generated at: {docx_path}")
    
    # 4. Save separate Excel Answer Key
    excel_key_path = os.path.splitext(docx_path)[0] + "_AnswerKey.xlsx"
    save_excel_answer_key(questions, excel_key_path, booklet_no)

if __name__ == "__main__":
    import sys
    excel_p = r"R:\Final_PYQ_\RAJASTHAN RAS EXAM PYQ.xlsx"
    docx_p = r"R:\Final_PYQ_\RAJASTHAN RAS EXAM PYQ.docx"
    topic_f = None
    
    if len(sys.argv) > 1:
        excel_p = sys.argv[1]
    if len(sys.argv) > 2:
        docx_p = sys.argv[2]
    if len(sys.argv) > 3:
        topic_f = sys.argv[3]
        
    generate_test_from_excel(excel_p, docx_p, topic_filter=topic_f)
